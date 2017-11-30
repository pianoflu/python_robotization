#encoding=UTF-8
from Zealink import zealinkmysql
from Zealink.public import *
from docx import Document
if __name__=="__main__":
    document = Document()
    conn = zealinkmysql.connect("AILI")
    sql = "select table_name,TABLE_COMMENT  from information_schema.tables where table_schema='xdb' and table_type='base table'"
    rds = sqlread(conn,sql)

    for item in rds:
        colInfoSql = "select table_name,column_name,column_type,is_nullable,column_key,COLUMN_COMMENT from INFORMATION_SCHEMA.Columns where table_name='%s' and table_schema='xdb'"
        colInfoSql = colInfoSql % item[0]
        colInfoRds = sqlread(conn,colInfoSql)
        p = document.add_paragraph("%s(%s)" % (item[1] if item[1] else "NULL",item[0]))
        table = document.add_table(rows=1, cols=5,style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'NAME'
        hdr_cells[1].text = 'TYPE'
        hdr_cells[2].text = 'NULLABLE'
        hdr_cells[3].text = 'PRIMARY'
        hdr_cells[4].text = 'COMMENTS'
        for detailItem in colInfoRds:
            print detailItem
            row_cells = table.add_row().cells
            row_cells[0].text = detailItem[1]
            row_cells[1].text = detailItem[2]
            row_cells[2].text = 'Y' if detailItem[3]=='YES' else ""
            row_cells[3].text = 'Y' if detailItem[4]=='PRI' else ""
            row_cells[4].text = detailItem[5]
        document.add_paragraph("")
    document.save('result.docx')



